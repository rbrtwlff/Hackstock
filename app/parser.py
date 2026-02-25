from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document


DOC_ORDER = ["Klage", "Klageerwiderung", "Replik", "Duplik", "Stellungnahme"]
HEADING_RE = re.compile(r"^(?:[IVXLC]+\.?|\d+(?:\.\d+)*\.?|[a-zA-Z]\)|\(\d+\))\s+")
BODY_ANCHOR_RE = re.compile(
    r"^(?:\s*(?:I|II|III|IV|V|VI|VII|VIII|IX|X)\.\s+|\s*\d+\.\s+|\s*\d+(?:\.\d+)+\s+|\s*[a-zA-Z]\)\s+|\s*\(\d+\)\s+|\s*(?:Sachverhalt|Antrag|Anträge|Begründung|Rechtliche Würdigung|I\.|II\.)\b)",
    re.IGNORECASE,
)
LAW_FIRM_RE = re.compile(
    r"(?:\b\d{5}\b|\b(?:straße|str\.|weg|platz|allee|gasse|ring)\b.*\d+|\b(?:Tel\.|Telefon|Fax|E-Mail|Email|www\.)\b|\b(?:Rechtsanwalt|Rechtsanwälte|Rechtsanwaltskanzlei|Kanzlei|Notar)\b|(?:>.*>\s*){2,}|(?:\|.*\|\s*){1,}|(?:.*,){3,})",
    re.IGNORECASE,
)
INTRO_RE = re.compile(
    r"\b(?:Gemäß|Nach|Wie|Dazu|Hierzu|vgl\.|Beschluss|Urteil|führt .* aus|heißt es)\b",
    re.IGNORECASE,
)
QUOTE_START_RE = re.compile(r'^\s*[„"»]|^\s*>\s+')
QUOTE_CLOSE_RE = re.compile(r'[“"«][^“"«]{0,20}$')
LEGAL_CITATION_RE = re.compile(
    r"(?:\b(?:BGH|OLG|LG|AG|BVerfG|EuGH)\b|\b[IVX]{1,5}\s*[A-Z]{1,3}\s*\d+/\d{2,4}\b|\bRdn\.?\s*\d+\b|\bRn\.?\s*\d+\b|\bBT-?Drucks\.?\s*\d+/\d+\b|§\s*\d+[a-zA-Z]*\s*(?:Abs\.?\s*\d+)?\s*(?:S\.?\s*\d+)?\s*[A-ZÄÖÜa-zäöü]{2,})"
)
CONNECTOR_RE = re.compile(r"^(?:und|oder|dass|weil|sodass|insbesondere|ferner)\b", re.IGNORECASE)


@dataclass
class ParsedParagraph:
    para_index: int
    text: str
    style: str
    is_heading: bool
    hierarchy_path: str
    continuation_group: str


@dataclass
class SemanticBlock:
    block_index: int
    block_type: str
    hierarchy_path: str
    text_original: str
    text_normalized: str
    source_paragraph_indexes: list[int]
    intro_text: str | None = None
    quote_text: str | None = None


def file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def table_to_json(table) -> tuple[str, str]:
    rows = []
    html_rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
        html_rows.append("<tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>")
    return json.dumps(rows, ensure_ascii=False), "<table>" + "".join(html_rows) + "</table>"


def is_heading(text: str, style_name: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if style_name.lower().startswith("heading"):
        return True
    if len(t) <= 90 and (HEADING_RE.match(t) or t.isupper()):
        return True
    return False


def is_body_anchor(text: str) -> bool:
    return bool(BODY_ANCHOR_RE.match(text.strip()))


def matches_law_firm_line(text: str) -> bool:
    return bool(LAW_FIRM_RE.search(text.strip()))


def should_ocr_normalize(paragraphs: list[str]) -> bool:
    if not paragraphs:
        return False
    short = [p for p in paragraphs if 0 < len(p.strip()) <= 35]
    return (len(short) / max(len(paragraphs), 1)) > 0.55


def normalize_ocr_lines(paragraphs: list[str]) -> list[str]:
    result = []
    buffer = ""
    for p in paragraphs:
        text = p.strip()
        if not text:
            if buffer:
                result.append(buffer.strip())
                buffer = ""
            continue
        if buffer and buffer.endswith("-"):
            buffer = buffer[:-1] + text
        elif buffer and len(buffer) < 180 and not HEADING_RE.match(text):
            buffer += " " + text
        else:
            if buffer:
                result.append(buffer.strip())
            buffer = text
    if buffer:
        result.append(buffer.strip())
    return result


def parse_docx(path: Path) -> tuple[list[str], list[tuple[str, str]]]:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    tables = [table_to_json(t) for t in doc.tables]
    return paragraphs, tables


def build_hierarchy(paragraphs: list[str], styles: list[str]) -> list[ParsedParagraph]:
    stack: list[str] = []
    parsed: list[ParsedParagraph] = []
    group = 0
    for i, (text, style) in enumerate(zip(paragraphs, styles)):
        head = is_heading(text, style)
        if head:
            group += 1
            token = text.strip()[:70]
            stack = stack[:3]
            stack.append(token)
        hierarchy = " > ".join(stack) if stack else "ROOT"
        parsed.append(
            ParsedParagraph(
                para_index=i,
                text=text.strip(),
                style=style,
                is_heading=head,
                hierarchy_path=hierarchy,
                continuation_group=f"g{group or 1}",
            )
        )
    return parsed


def _quote_like(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    starts_lower = stripped[0].islower()
    long_and_legal = len(stripped) > 80 and bool(LEGAL_CITATION_RE.search(stripped))
    quote_punct = stripped.startswith(("„", '"')) or stripped.count(";") >= 2 or stripped.count("(") >= 2
    return starts_lower or long_and_legal or quote_punct


def _is_subheading(para: ParsedParagraph) -> bool:
    return para.is_heading or (len(para.text) < 90 and para.style.lower().startswith("heading"))


def _is_short_merge_candidate(prev_text: str, text: str) -> bool:
    if len(text) >= 40 and len(text.split()) >= 7:
        return False
    if prev_text.rstrip().endswith((".", "?", "!", ":")):
        return False
    stripped = text.strip()
    return bool(stripped and (stripped[0].islower() or CONNECTOR_RE.match(stripped)))


def build_semantic_blocks(paragraphs: list[ParsedParagraph]) -> list[SemanticBlock]:
    blocks: list[SemanticBlock] = []
    rubrum_parts: list[ParsedParagraph] = []
    i = 0
    while i < len(paragraphs) and not is_body_anchor(paragraphs[i].text):
        rubrum_parts.append(paragraphs[i])
        i += 1
    if rubrum_parts:
        rubrum_text = "\n".join(p.text for p in rubrum_parts)
        blocks.append(
            SemanticBlock(
                block_index=len(blocks),
                block_type="RUBRUM_META",
                hierarchy_path="RUBRUM",
                text_original=rubrum_text,
                text_normalized=" ".join(rubrum_text.split()),
                source_paragraph_indexes=[p.para_index for p in rubrum_parts],
            )
        )

    pending_heading: ParsedParagraph | None = None
    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()

        if len(text) < 120 and matches_law_firm_line(text) and not is_body_anchor(text) and not para.is_heading:
            if blocks and blocks[0].block_type == "RUBRUM_META":
                merged = blocks[0].text_original + "\n" + text
                blocks[0].text_original = merged
                blocks[0].text_normalized = " ".join(merged.split())
                blocks[0].source_paragraph_indexes.append(para.para_index)
            i += 1
            continue

        if _is_subheading(para):
            pending_heading = para
            i += 1
            if i < len(paragraphs) and _is_subheading(paragraphs[i]):
                blocks.append(
                    SemanticBlock(
                        block_index=len(blocks),
                        block_type="BODY",
                        hierarchy_path=para.hierarchy_path,
                        text_original=para.text,
                        text_normalized=" ".join(para.text.split()),
                        source_paragraph_indexes=[para.para_index],
                    )
                )
                pending_heading = None
            continue

        quote_start = bool(QUOTE_START_RE.match(text)) or (LEGAL_CITATION_RE.search(text) and len(text) > 80)
        intro_para: ParsedParagraph | None = None
        if quote_start and i > 0:
            prev = paragraphs[i - 1]
            prev_text = prev.text.strip()
            if not prev.is_heading and INTRO_RE.search(prev_text) and prev_text.endswith(":"):
                intro_para = prev
                if blocks and intro_para.para_index in blocks[-1].source_paragraph_indexes:
                    blocks.pop()

        if quote_start:
            quote_items = [para]
            j = i + 1
            while j < len(paragraphs):
                cand = paragraphs[j]
                cand_text = cand.text.strip()
                if is_body_anchor(cand_text) or cand.is_heading:
                    break
                if _quote_like(cand_text):
                    quote_items.append(cand)
                    if QUOTE_CLOSE_RE.search(cand_text):
                        if j + 1 < len(paragraphs) and not _quote_like(paragraphs[j + 1].text):
                            j += 1
                            break
                    j += 1
                    continue
                break
            quote_text = "\n".join(p.text for p in quote_items)
            source_ids = [p.para_index for p in quote_items]
            block_type = "QUOTE_BLOCK"
            original_text = quote_text
            hierarchy = para.hierarchy_path
            if intro_para:
                block_type = "BODY_WITH_QUOTE"
                original_text = intro_para.text + "\n" + quote_text
                source_ids = [intro_para.para_index] + source_ids
                hierarchy = intro_para.hierarchy_path
            blocks.append(
                SemanticBlock(
                    block_index=len(blocks),
                    block_type=block_type,
                    hierarchy_path=hierarchy,
                    text_original=original_text,
                    text_normalized=" ".join(original_text.split()),
                    source_paragraph_indexes=source_ids,
                    intro_text=intro_para.text if intro_para else None,
                    quote_text=quote_text,
                )
            )
            i = j
            continue

        texts = [text]
        source = [para.para_index]
        hierarchy = para.hierarchy_path
        if pending_heading:
            texts.insert(0, pending_heading.text)
            source.insert(0, pending_heading.para_index)
            hierarchy = pending_heading.hierarchy_path
            pending_heading = None
        j = i + 1
        while j < len(paragraphs):
            cand = paragraphs[j]
            if _is_subheading(cand) or is_body_anchor(cand.text):
                break
            if _is_short_merge_candidate(texts[-1], cand.text):
                texts.append(cand.text)
                source.append(cand.para_index)
                j += 1
                continue
            break
        merged = "\n".join(texts)
        blocks.append(
            SemanticBlock(
                block_index=len(blocks),
                block_type="BODY",
                hierarchy_path=hierarchy,
                text_original=merged,
                text_normalized=" ".join(merged.split()),
                source_paragraph_indexes=source,
            )
        )
        i = j

    return blocks
