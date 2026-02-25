from __future__ import annotations

import asyncio
import csv
import hashlib
import json
import logging
from pathlib import Path

from docx import Document

from app.config import AppConfig
from app.db import db_conn
from app.llm import LLMClient
from app.parser import DOC_ORDER, build_hierarchy, build_semantic_blocks, file_hash, normalize_ocr_lines, should_ocr_normalize, table_to_json

logger = logging.getLogger("pipeline")


class Pipeline:
    def __init__(self, config: AppConfig):
        self.config = config
        self.llm = LLMClient(config)

    def run_all(self):
        self.import_documents()
        self.normalize_documents()
        asyncio.run(self.analyze_semantic_blocks())
        self.build_arguments()
        asyncio.run(self.propose_links())

    def retry_failed(self):
        asyncio.run(self.analyze_semantic_blocks(only_failed=True))
        asyncio.run(self.propose_links(only_failed=True))

    def import_documents(self):
        manifest = Path("data/manifest.csv")
        inbox = Path("data/inbox")
        rows = list(csv.DictReader(manifest.read_text(encoding="utf-8-sig").splitlines()))
        rows.sort(key=lambda r: (DOC_ORDER.index(r["doc_type"]), r["date"]))
        with db_conn(self.config.db_path) as conn:
            for row in rows:
                p = inbox / row["filename"]
                if not p.exists():
                    logger.error("Missing file %s", p)
                    continue
                ch = file_hash(p)
                conn.execute(
                    """INSERT INTO documents(doc_id,side,doc_type,date,filename,content_hash)
                    VALUES(?,?,?,?,?,?)
                    ON CONFLICT(doc_id) DO UPDATE SET side=excluded.side, doc_type=excluded.doc_type, date=excluded.date,
                    filename=excluded.filename, content_hash=excluded.content_hash""",
                    (row["doc_id"], row["side"], row["doc_type"], row["date"], row["filename"], ch),
                )
                document_id = conn.execute("SELECT id FROM documents WHERE doc_id=?", (row["doc_id"],)).fetchone()[0]
                self._ingest_doc(conn, document_id, p)

    def _ingest_doc(self, conn, document_id: int, path: Path):
        doc = Document(path)
        paras, styles = [], []
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if not txt:
                continue
            paras.append(txt)
            styles.append(getattr(p.style, "name", "Normal"))

        conn.execute("DELETE FROM table_blocks WHERE document_id=?", (document_id,))
        conn.execute("DELETE FROM semantic_block_sources WHERE block_id IN (SELECT id FROM semantic_blocks WHERE document_id=?)", (document_id,))
        conn.execute("DELETE FROM semantic_block_jobs WHERE block_id IN (SELECT id FROM semantic_blocks WHERE document_id=?)", (document_id,))
        conn.execute("DELETE FROM semantic_block_analysis WHERE block_id IN (SELECT id FROM semantic_blocks WHERE document_id=?)", (document_id,))
        conn.execute("DELETE FROM semantic_blocks WHERE document_id=?", (document_id,))
        conn.execute("DELETE FROM paragraphs WHERE document_id=?", (document_id,))

        parsed = build_hierarchy(paras, styles)
        para_ids_by_idx = {}
        for para in parsed:
            h = hashlib.sha256(f"{document_id}:{para.para_index}:{para.text}".encode()).hexdigest()
            conn.execute(
                """INSERT INTO paragraphs(document_id,para_index,text,style,is_heading,hierarchy_path,continuation_group,content_hash)
                VALUES(?,?,?,?,?,?,?,?)""",
                (document_id, para.para_index, para.text, para.style, int(para.is_heading), para.hierarchy_path, para.continuation_group, h),
            )
            para_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
            para_ids_by_idx[para.para_index] = para_id

        semantic_blocks = build_semantic_blocks(parsed)
        for b in semantic_blocks:
            conn.execute(
                """INSERT INTO semantic_blocks(document_id,block_index,block_type,hierarchy_path,text_original,text_normalized,intro_text,quote_text)
                VALUES(?,?,?,?,?,?,?,?)""",
                (document_id, b.block_index, b.block_type, b.hierarchy_path, b.text_original, b.text_normalized, b.intro_text, b.quote_text),
            )
            block_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
            for ord_idx, para_idx in enumerate(b.source_paragraph_indexes):
                para_id = para_ids_by_idx.get(para_idx)
                if para_id:
                    conn.execute(
                        "INSERT OR IGNORE INTO semantic_block_sources(block_id,paragraph_id,ord) VALUES(?,?,?)",
                        (block_id, para_id, ord_idx),
                    )
            conn.execute("INSERT OR IGNORE INTO semantic_block_jobs(block_id,status) VALUES(?, 'PENDING')", (block_id,))

        for idx, table in enumerate(doc.tables):
            cells_json, render_html = table_to_json(table)
            conn.execute(
                """INSERT INTO table_blocks(document_id,block_index,cells_json,render_html)
                VALUES(?,?,?,?)""",
                (document_id, idx, cells_json, render_html),
            )

        conn.execute(
            "UPDATE documents SET raw_paragraph_count=?, semantic_block_count=? WHERE id=?",
            (len(parsed), len(semantic_blocks)),
            document_id,
        )

    def normalize_documents(self):
        with db_conn(self.config.db_path) as conn:
            docs = conn.execute("SELECT id FROM documents").fetchall()
            for d in docs:
                rows = conn.execute("SELECT id,text FROM paragraphs WHERE document_id=? ORDER BY para_index", (d[0],)).fetchall()
                texts = [r["text"] for r in rows]
                trig = should_ocr_normalize(texts)
                if trig:
                    normalized = normalize_ocr_lines(texts)
                    for i, t in enumerate(normalized):
                        if i < len(rows):
                            conn.execute("UPDATE paragraphs SET text=? WHERE id=?", (t, rows[i]["id"]))
                conn.execute("UPDATE documents SET ocr_normalized=? WHERE id=?", (1 if trig else 0, d[0]))

    async def analyze_semantic_blocks(self, only_failed: bool = False):
        with db_conn(self.config.db_path) as conn:
            query = f"""SELECT sb.id,sb.text_normalized,sb.hierarchy_path,sbj.attempts
            FROM semantic_blocks sb JOIN semantic_block_jobs sbj ON sb.id=sbj.block_id
            WHERE sb.block_type IN ('BODY','BODY_WITH_QUOTE','QUOTE_BLOCK')
            AND sbj.status IN ({"'FAILED'" if only_failed else "'PENDING','FAILED'"})"""
            jobs = conn.execute(query).fetchall()

        for job in jobs:
            block_id, text, hierarchy, attempts = job
            with db_conn(self.config.db_path) as conn:
                conn.execute("UPDATE semantic_block_jobs SET status='RUNNING',attempts=attempts+1,updated_at=CURRENT_TIMESTAMP WHERE block_id=?", (block_id,))
            try:
                prev_next = self._neighbor_context(block_id)
                result = await self.llm.analyze_paragraph(text, f"{hierarchy}\n{prev_next}")
                with db_conn(self.config.db_path) as conn:
                    canonical_issues = [self._canonicalize_issue(conn, x) for x in result.issues]
                    conn.execute(
                        """INSERT INTO semantic_block_analysis(block_id,keywords_json,issues_json,role,summary_3_sentences,
                        continuation_of_previous,continuation_reason,citations_norms_json,citations_cases_json,citations_contract_json,citations_exhibits_json)
                        VALUES(?,?,?,?,?,?,?,?,?,?,?)
                        ON CONFLICT(block_id) DO UPDATE SET keywords_json=excluded.keywords_json, issues_json=excluded.issues_json,
                        role=excluded.role, summary_3_sentences=excluded.summary_3_sentences,
                        continuation_of_previous=excluded.continuation_of_previous, continuation_reason=excluded.continuation_reason,
                        citations_norms_json=excluded.citations_norms_json,citations_cases_json=excluded.citations_cases_json,
                        citations_contract_json=excluded.citations_contract_json,citations_exhibits_json=excluded.citations_exhibits_json""",
                        (
                            block_id,
                            json.dumps(result.keywords, ensure_ascii=False),
                            json.dumps(canonical_issues, ensure_ascii=False),
                            result.role.value,
                            result.summary_3_sentences,
                            int(result.continuation_of_previous),
                            result.continuation_reason,
                            json.dumps(result.citations_norms, ensure_ascii=False),
                            json.dumps(result.citations_cases, ensure_ascii=False),
                            json.dumps(result.citations_contract, ensure_ascii=False),
                            json.dumps(result.citations_exhibits, ensure_ascii=False),
                        ),
                    )
                    conn.execute("UPDATE semantic_block_jobs SET status='DONE',last_error=NULL,updated_at=CURRENT_TIMESTAMP WHERE block_id=?", (block_id,))
            except Exception as exc:
                with db_conn(self.config.db_path) as conn:
                    status = "FAILED" if attempts + 1 >= 3 else "PENDING"
                    conn.execute("UPDATE semantic_block_jobs SET status=?,last_error=?,updated_at=CURRENT_TIMESTAMP WHERE block_id=?", (status, str(exc)[:500], block_id))

    def _neighbor_context(self, block_id: int) -> str:
        with db_conn(self.config.db_path) as conn:
            row = conn.execute("SELECT document_id,block_index FROM semantic_blocks WHERE id=?", (block_id,)).fetchone()
            prev = conn.execute("SELECT text_normalized FROM semantic_blocks WHERE document_id=? AND block_index=?", (row[0], row[1] - 1)).fetchone()
            nxt = conn.execute("SELECT text_normalized FROM semantic_blocks WHERE document_id=? AND block_index=?", (row[0], row[1] + 1)).fetchone()
            return f"Vorher: {(prev['text_normalized'] if prev else '')}\nNachher: {(nxt['text_normalized'] if nxt else '')}"

    def _canonicalize_issue(self, conn, issue: str) -> str:
        norm = issue.strip().lower()
        rows = conn.execute("SELECT canonical,synonyms_json FROM issue_vocab").fetchall()
        for r in rows:
            canonical = r["canonical"]
            syns = json.loads(r["synonyms_json"])
            if norm == canonical.lower() or norm in [s.lower() for s in syns]:
                return canonical
            if canonical.lower() in norm or norm in canonical.lower():
                return canonical
        canonical = issue.strip()
        conn.execute("INSERT OR IGNORE INTO issue_vocab(canonical,synonyms_json) VALUES(?,?)", (canonical, json.dumps([], ensure_ascii=False)))
        return canonical

    def build_arguments(self):
        with db_conn(self.config.db_path) as conn:
            conn.execute("DELETE FROM semantic_block_arguments")
            conn.execute("DELETE FROM arguments")
            rows = conn.execute(
                """SELECT sb.id,sb.document_id,d.side,sb.hierarchy_path,sb.block_type,sb.text_original
                FROM semantic_blocks sb JOIN documents d ON sb.document_id=d.id
                WHERE sb.block_type IN ('BODY','BODY_WITH_QUOTE','QUOTE_BLOCK')
                ORDER BY sb.document_id,sb.block_index"""
            ).fetchall()
            cache = {}
            for r in rows:
                key = (r["document_id"], r["side"], r["hierarchy_path"], r["block_type"])
                if key not in cache:
                    title = f"{r['hierarchy_path']} / {r['block_type']}"
                    conn.execute("INSERT INTO arguments(document_id,side,title,hierarchy_path) VALUES(?,?,?,?)", (r["document_id"], r["side"], title, r["hierarchy_path"]))
                    cache[key] = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
                conn.execute("INSERT OR IGNORE INTO semantic_block_arguments(block_id,argument_id) VALUES(?,?)", (r["id"], cache[key]))

    async def propose_links(self, only_failed: bool = False):
        with db_conn(self.config.db_path) as conn:
            plaintiff = conn.execute("SELECT a.id,a.title FROM arguments a WHERE side='PLAINTIFF'").fetchall()
            defendant = conn.execute("SELECT a.id,a.title FROM arguments a WHERE side='DEFENDANT'").fetchall()
            for p in plaintiff:
                for d in defendant:
                    conn.execute("INSERT OR IGNORE INTO links_jobs(from_argument_id,to_argument_id,status) VALUES(?,?,'PENDING')", (p[0], d[0]))
            status_filter = "('FAILED')" if only_failed else "('PENDING','FAILED')"
            jobs = conn.execute(f"SELECT id,from_argument_id,to_argument_id,attempts FROM links_jobs WHERE status IN {status_filter}").fetchall()
        for j in jobs:
            with db_conn(self.config.db_path) as conn:
                conn.execute("UPDATE links_jobs SET status='RUNNING',attempts=attempts+1 WHERE id=?", (j["id"],))
                p = conn.execute("SELECT title FROM arguments WHERE id=?", (j["from_argument_id"],)).fetchone()
                d = conn.execute("SELECT title FROM arguments WHERE id=?", (j["to_argument_id"],)).fetchone()
            try:
                proposal = await self.llm.classify_link(p[0], d[0])
                with db_conn(self.config.db_path) as conn:
                    conn.execute(
                        """INSERT INTO links(from_argument_id,to_argument_id,link_type,confidence,rationale_short,status)
                        VALUES(?,?,?,?,?,'proposed')""",
                        (j["from_argument_id"], j["to_argument_id"], proposal.link_type.value, proposal.confidence, proposal.rationale_short),
                    )
                    conn.execute("UPDATE links_jobs SET status='DONE',last_error=NULL WHERE id=?", (j["id"],))
            except Exception as exc:
                with db_conn(self.config.db_path) as conn:
                    status = "FAILED" if j["attempts"] + 1 >= 3 else "PENDING"
                    conn.execute("UPDATE links_jobs SET status=?,last_error=? WHERE id=?", (status, str(exc)[:500], j["id"]))
